#!/usr/bin/env python3
"""
WhaTap Documentation Word Export with Cover Page and Styling
PDF와 동일한 스타일을 Word 문서에 적용
"""

import os
import sys
import subprocess
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("❌ python-docx 라이브러리가 필요합니다.")
    print("설치: pip install python-docx")
    sys.exit(1)

# 제품 정보
PRODUCT_INFO = {
    'java': {'title': 'Java Monitoring', 'category': 'APM'},
    'nodejs': {'title': 'Node.js Monitoring', 'category': 'APM'},
    'python': {'title': 'Python Monitoring', 'category': 'APM'},
    'php': {'title': 'PHP Monitoring', 'category': 'APM'},
    'dotnet': {'title': '.NET Monitoring', 'category': 'APM'},
    'golang': {'title': 'Go Monitoring', 'category': 'APM'},
    
    'mysql': {'title': 'MySQL Monitoring', 'category': 'Database'},
    'postgresql': {'title': 'PostgreSQL Monitoring', 'category': 'Database'},
    'mongodb': {'title': 'MongoDB Monitoring', 'category': 'Database'},
    'oracle': {'title': 'Oracle Monitoring', 'category': 'Database'},
    'oracle-pro': {'title': 'Oracle Pro Monitoring', 'category': 'Database'},
    'mssql': {'title': 'MS SQL Monitoring', 'category': 'Database'},
    'redis': {'title': 'Redis Monitoring', 'category': 'Database'},
    'tibero': {'title': 'Tibero Monitoring', 'category': 'Database'},
    'altibase': {'title': 'Altibase Monitoring', 'category': 'Database'},
    'cubrid': {'title': 'Cubrid Monitoring', 'category': 'Database'},
    'db2': {'title': 'DB2 Monitoring', 'category': 'Database'},
    'sapase': {'title': 'SAP ASE Monitoring', 'category': 'Database'},
    
    'kubernetes': {'title': 'Kubernetes Monitoring', 'category': 'Infrastructure'},
    'server': {'title': 'Server Monitoring', 'category': 'Infrastructure'},
    'npm': {'title': 'Network Performance Monitoring', 'category': 'Infrastructure'},
}

def create_cover_page(doc, product_name, cover_bg_path):
    """커버 페이지 생성 - 이미지 위에 텍스트 오버레이"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    info = PRODUCT_INFO.get(product_name, {'title': product_name.title(), 'category': 'Monitoring'})
    
    # 커버 페이지 섹션 - 여백 제거
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 height
    section.page_width = Inches(8.27)    # A4 width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # 배경 이미지를 페이지에 고정 (Behind Text)
    if os.path.exists(cover_bg_path):
        try:
            # 이미지 삽입
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            pic = run.add_picture(cover_bg_path, width=Inches(7.27))
            
            # 이미지를 Behind Text로 설정
            inline = run._element
            docPr = inline.xpath('.//wp:docPr')[0]
            
            # anchor 요소 생성 (floating)
            anchor = OxmlElement('wp:anchor')
            
            # 기본 속성 복사
            for child in inline:
                anchor.append(child)
            
            # Behind text 설정
            anchor.set(qn('behindDoc'), '1')
            anchor.set(qn('locked'), '0')
            anchor.set(qn('layoutInCell'), '1')
            anchor.set(qn('allowOverlap'), '1')
            
            # Position 설정 (페이지 중앙 하단)
            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'page')
            positionH_align = OxmlElement('wp:align')
            positionH_align.text = 'center'
            positionH.append(positionH_align)
            anchor.append(positionH)
            
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'page')
            posV_offset = OxmlElement('wp:posOffset')
            posV_offset.text = '4000000'  # 위치 조정 (EMU 단위)
            positionV.append(posV_offset)
            anchor.append(positionV)
            
            # 원본 inline을 anchor로 교체
            inline.getparent().replace(inline, anchor)
            
        except Exception as e:
            print(f"⚠️  배경 이미지 오버레이 실패: {e}")
            print(f"    일반 이미지로 추가합니다.")
    
    # 상단 여백 (텍스트 위치 조정)
    for _ in range(10):
        doc.add_paragraph()
    
    # 제목 (이미지 위에)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(info['title'])
    title_run.font.size = Pt(40)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(41, 108, 242)  # WhaTap Blue
    title_run.font.name = 'Noto Sans KR'
    
    # 공백
    for _ in range(3):
        doc.add_paragraph()
    
    # 릴리즈 날짜
    today = datetime.now()
    release_date = f"release date. {today.year}. {today.month:02d}. {today.day:02d}."
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(release_date)
    date_run.font.size = Pt(16)
    date_run.font.color.rgb = RGBColor(102, 102, 102)
    date_run.font.name = 'Noto Sans KR'
    
    # 페이지 구분
    doc.add_page_break()

def apply_whatap_styles(doc):
    """WhaTap 문서 스타일 적용 (PDF CSS 기반)"""
    
    # 기본 스타일
    styles = doc.styles
    
    # Heading 1 스타일
    h1 = styles['Heading 1']
    h1.font.size = Pt(40)  # 2.5rem
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(34, 34, 34)
    h1.font.name = 'Noto Sans KR'
    
    # Heading 2 스타일
    h2 = styles['Heading 2']
    h2.font.size = Pt(17.6)  # 1.1rem
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(34, 34, 34)
    h2.font.name = 'Noto Sans KR'
    
    # Heading 3 스타일
    h3 = styles['Heading 3']
    h3.font.size = Pt(20)  # 1.25rem
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(34, 34, 34)
    h3.font.name = 'Noto Sans KR'
    
    # Normal 스타일
    normal = styles['Normal']
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(34, 34, 34)
    normal.font.name = 'Noto Sans KR'

def convert_with_pandoc(product_name, docs_dir='./docs', output_dir='./exports/word'):
    """Pandoc으로 마크다운 → Word 변환"""
    
    product_dir = Path(docs_dir) / product_name
    
    if not product_dir.exists():
        print(f"⚠️  {product_name}: 폴더가 존재하지 않습니다.")
        return None
    
    # 마크다운 파일 수집
    md_files = sorted(product_dir.rglob('*.md')) + sorted(product_dir.rglob('*.mdx'))
    
    if not md_files:
        print(f"⚠️  {product_name}: 마크다운 파일이 없습니다.")
        return None
    
    output_path = Path(output_dir) / f"{product_name}-manual.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    print(f"📝 {product_name}: {len(md_files)}개 파일 변환 중...")
    
    # Pandoc 명령어
    file_list = ' '.join([f'"{f}"' for f in md_files])
    command = f'''pandoc {file_list} -o "{output_path}" \
        --toc \
        --toc-depth=3 \
        -s \
        --number-sections \
        -V documentclass=article \
        -V papersize=a4 \
        -V fontsize=10pt \
        -V geometry:margin=25mm \
        -V mainfont="Noto Sans KR" \
        -V CJKmainfont="Noto Sans KR"'''
    
    try:
        subprocess.run(command, shell=True, check=True, capture_output=True)
        print(f"✅ {product_name}: {output_path} 생성 완료 (Pandoc)")
        return output_path
    except subprocess.CalledProcessError as e:
        print(f"❌ {product_name}: Pandoc 변환 실패")
        print(e.stderr.decode())
        return None

def add_cover_to_word(docx_path, product_name, cover_bg_path='./static/img/cover-background.png'):
    """기존 Word 문서에 커버 페이지 추가"""
    
    try:
        doc = Document(str(docx_path))
        
        # 기존 내용을 임시 저장
        temp_path = docx_path.with_suffix('.tmp.docx')
        doc.save(str(temp_path))
        
        # 새 문서 생성 (커버 먼저)
        new_doc = Document()
        create_cover_page(new_doc, product_name, cover_bg_path)
        apply_whatap_styles(new_doc)
        
        # 기존 내용 병합
        temp_doc = Document(str(temp_path))
        for element in temp_doc.element.body:
            new_doc.element.body.append(element)
        
        # 저장
        new_doc.save(str(docx_path))
        
        # 임시 파일 삭제
        temp_path.unlink()
        
        print(f"✅ {product_name}: 커버 페이지 추가 완료")
        
    except Exception as e:
        print(f"❌ {product_name}: 커버 추가 실패 - {e}")

def main():
    """메인 실행"""
    print("=" * 60)
    print("📚 WhaTap Documentation Export to Word (with Cover)")
    print("=" * 60)
    
    # 명령줄 인자
    products = sys.argv[1:] if len(sys.argv) > 1 else list(PRODUCT_INFO.keys())
    
    if not products:
        print("사용법: python3 export-word-styled.py [제품명...]")
        print(f"예시: python3 export-word-styled.py java mysql")
        return
    
    print(f"📦 변환 대상: {', '.join(products)}\n")
    
    for product in products:
        # 1. Pandoc으로 기본 Word 생성
        docx_path = convert_with_pandoc(product)
        
        if docx_path:
            # 2. 커버 페이지 추가 및 스타일 적용
            add_cover_to_word(docx_path, product)
    
    print("\n" + "=" * 60)
    print("✨ 변환 완료!")
    print(f"📂 출력 위치: ./exports/word/")
    print("=" * 60)

if __name__ == '__main__':
    main()